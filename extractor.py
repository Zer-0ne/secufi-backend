"""
🔓 Enhanced Universal Document Text Extraction System with Table Support
Supports: PDF (including password-protected), Scanned PDF, Images, DOCX, CSV, Excel, and more

Author: AI Assistant
Version: 3.0 with Enhanced Table Extraction
"""

import os
import sys
import argparse
from pathlib import Path
from typing import Optional, Dict, List
import mimetypes
import io
import re

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# PDF Extraction
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

# Enhanced table extraction
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import tabula
    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False

# OCR for scanned PDFs and images
try:
    from PIL import Image
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# DOCX extraction
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel extraction
try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# CSV extraction
import csv


class DocumentExtractor:
    """Universal document text extraction class with enhanced table support"""
    
    def __init__(self, file_path: str, password: Optional[str] = None):
        """
        Initialize extractor with file path and optional password
        
        Args:
            file_path (str): Path to the document file
            password (str, optional): Password for encrypted/protected documents
            
        Raises:
            FileNotFoundError: If file doesn't exist
        """
        self.file_path = Path(file_path)
        self.password = password
        
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        self.mime_type, _ = mimetypes.guess_type(str(self.file_path))
        self.extension = self.file_path.suffix.lower()
        
    def extract(self) -> Dict[str, any]:
        """
        Extract text from document based on file type
        
        Returns:
            Dict: Dictionary containing extracted text and metadata
        """
        print(f"⟳ Processing: {self.file_path.name}")
        print(f"⟳ Type: {self.mime_type or self.extension}")
        
        if self.password:
            print(f"⟳ Password protected mode enabled")
        
        # Route to appropriate extractor
        if self.extension == '.pdf':
            return self._extract_pdf()
        elif self.extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']:
            return self._extract_image()
        elif self.extension in ['.docx', '.doc']:
            return self._extract_docx()
        elif self.extension == '.csv':
            return self._extract_csv()
        elif self.extension in ['.xlsx', '.xls']:
            return self._extract_excel()
        elif self.extension == '.txt':
            return self._extract_text()
        else:
            return self._extract_generic()
    
    def _extract_pdf(self) -> Dict[str, any]:
        """
        Extract text from PDF with enhanced table extraction
        
        Returns:
            Dict: Extraction result with text and metadata
        """
        text = ""
        method_used = None
        has_tables = False
        
        # Try pdfplumber first for best table extraction
        if PDFPLUMBER_AVAILABLE:
            try:
                print("⟳ Using pdfplumber for table-aware extraction...")
                
                pdf_config = {'password': self.password} if self.password else {}
                
                with pdfplumber.open(str(self.file_path), **pdf_config) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        print(f"  ⟳ Processing page {page_num}/{len(pdf.pages)}...")
                        
                        text += f"\n--- Page {page_num} ---\n"
                        
                        # Extract tables first
                        tables = page.extract_tables()
                        
                        if tables:
                            has_tables = True
                            for table_num, table in enumerate(tables, 1):
                                text += f"\n### Table {table_num} ###\n"
                                text += self._format_table(table)
                                text += "\n"
                        
                        # Extract remaining text (non-table text)
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    
                    method_used = "pdfplumber (table-aware)"
                    print("✓ PDF extracted successfully with pdfplumber")
                    
            except Exception as e:
                print(f"⚠ pdfplumber failed: {str(e)}")
                text = ""
        
        # Fallback to PyMuPDF if pdfplumber not available or failed
        if not text and PYMUPDF_AVAILABLE:
            try:
                print("⟳ Fallback to PyMuPDF...")
                doc = fitz.open(str(self.file_path))
                
                # Check if PDF is encrypted
                if doc.is_pdf and doc.is_encrypted:
                    print("⟳ PDF is encrypted/locked detected")
                    if not self.password:
                        return {
                            'text': "❌ PDF is password protected. Provide password with --password flag",
                            'method': 'Error',
                            'file_name': self.file_path.name,
                            'success': False
                        }
                    
                    auth_result = doc.authenticate(self.password)
                    if not auth_result:
                        return {
                            'text': "❌ Incorrect password provided",
                            'method': 'Error',
                            'file_name': self.file_path.name,
                            'success': False
                        }
                    print("✓ PDF unlocked successfully with PyMuPDF")
                
                # Extract text from pages
                for page_num, page in enumerate(doc, 1):
                    try:
                        page_text = page.get_text()
                        text += f"\n--- Page {page_num} ---\n{page_text}"
                    except Exception as page_error:
                        print(f"⚠ Error on page {page_num}: {page_error}")
                        text += f"\n--- Page {page_num} ---\n[Error extracting page content]"
                
                doc.close()
                method_used = "PyMuPDF"
                
                # If text is too short, PDF might be scanned
                if len(text.strip()) < 100:
                    print("⚠ Low text content detected, trying OCR...")
                    text = self._extract_pdf_with_ocr()
                    method_used = "PyMuPDF + OCR"
                    
            except Exception as e:
                print(f"⚠ PyMuPDF failed: {str(e)}")
                text = ""
        
        # Try tabula for table extraction if other methods failed
        if not text and TABULA_AVAILABLE:
            try:
                print("⟳ Using tabula for table extraction...")
                tables = tabula.read_pdf(str(self.file_path), 
                                        password=self.password,
                                        pages='all', 
                                        multiple_tables=True)
                
                text = ""
                for i, df in enumerate(tables, 1):
                    text += f"\n### Table {i} ###\n"
                    text += df.to_string(index=False)
                    text += "\n\n"
                
                method_used = "tabula-py"
                has_tables = True
                
            except Exception as e:
                print(f"⚠ tabula failed: {str(e)}")
        
        # Fallback to pypdf
        if not text and PYPDF_AVAILABLE:
            try:
                print("⟳ Fallback to pypdf...")
                reader = PdfReader(str(self.file_path))
                
                if reader.is_encrypted:
                    print("⟳ PDF is encrypted/locked detected (pypdf)")
                    if not self.password:
                        return {
                            'text': "❌ PDF is password protected. Provide password with --password flag",
                            'method': 'Error',
                            'file_name': self.file_path.name,
                            'success': False
                        }
                    
                    decrypt_result = reader.decrypt(self.password)
                    if not decrypt_result:
                        return {
                            'text': "❌ Incorrect password provided",
                            'method': 'Error',
                            'file_name': self.file_path.name,
                            'success': False
                        }
                    print("✓ PDF unlocked successfully with pypdf")
                
                for page_num, page in enumerate(reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        text += f"\n--- Page {page_num} ---\n{page_text}"
                    except Exception as page_error:
                        print(f"⚠ Error on page {page_num}: {page_error}")
                        text += f"\n--- Page {page_num} ---\n[Error extracting page content]"
                
                method_used = "pypdf"
                
            except Exception as e:
                print(f"⚠ pypdf failed: {str(e)}")
        
        # If still no text, use OCR
        if not text or len(text.strip()) < 50:
            if OCR_AVAILABLE:
                print("⟳ Using OCR for scanned PDF...")
                text = self._extract_pdf_with_ocr()
                method_used = "OCR (Tesseract)"
            else:
                text = "❌ OCR not available. Install: pip install pytesseract pdf2image pillow"
        
        return {
            'text': text.strip(),
            'method': method_used,
            'file_name': self.file_path.name,
            'file_size': f"{self.file_path.stat().st_size / 1024:.2f}KB",
            'page_count': self._estimate_pages(text),
            'char_count': len(text),
            'has_tables': has_tables,
            'success': len(text.strip()) > 0
        }
    
    def _format_table(self, table: List[List[str]]) -> str:
        """
        Format extracted table as markdown table
        
        Args:
            table: List of rows, where each row is a list of cells
            
        Returns:
            str: Formatted markdown table
        """
        if not table or len(table) == 0:
            return ""
        
        # Clean cells
        cleaned_table = []
        for row in table:
            cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
            cleaned_table.append(cleaned_row)
        
        # Find max column width
        max_cols = max(len(row) for row in cleaned_table)
        
        # Pad rows to have same column count
        for row in cleaned_table:
            while len(row) < max_cols:
                row.append("")
        
        # Format as markdown
        formatted = ""
        
        # Header row
        if cleaned_table:
            header = cleaned_table[0]
            formatted += "| " + " | ".join(header) + " |\n"
            formatted += "| " + " | ".join(['---'] * len(header)) + " |\n"
            
            # Data rows
            for row in cleaned_table[1:]:
                formatted += "| " + " | ".join(row) + " |\n"
        
        return formatted
    
    def _extract_pdf_with_ocr(self) -> str:
        """
        Extract text from PDF using OCR (Tesseract)
        
        Returns:
            str: Extracted text from PDF pages
        """
        if not OCR_AVAILABLE:
            return "❌ OCR libraries not installed. Install: pip install pytesseract pdf2image pillow"
        
        try:
            kwargs = {'dpi': 300}
            images = convert_from_path(str(self.file_path), **kwargs)
            text = ""
            
            for page_num, image in enumerate(images, 1):
                print(f"  ⟳ OCR processing page {page_num}/{len(images)}...")
                page_text = pytesseract.image_to_string(image, lang='eng')
                text += f"\n--- Page {page_num} ---\n{page_text}"
            
            print(f"✓ OCR completed: {len(images)} pages processed")
            return text
            
        except Exception as e:
            print(f"⚠ OCR failed: {str(e)}")
            return f"❌ OCR Error: {str(e)}"
    
    def _extract_image(self) -> Dict[str, any]:
        """
        Extract text from image using OCR
        
        Returns:
            Dict: Extraction result with text and metadata
        """
        if not OCR_AVAILABLE:
            return {
                'text': "❌ OCR not available. Install: pip install pytesseract pillow",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
        
        try:
            image = Image.open(str(self.file_path))
            
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            text = pytesseract.image_to_string(image, lang='eng')
            
            return {
                'text': text.strip(),
                'method': 'Tesseract OCR',
                'file_name': self.file_path.name,
                'file_size': f"{self.file_path.stat().st_size / 1024:.2f}KB",
                'image_size': f"{image.width}x{image.height}",
                'image_mode': image.mode,
                'char_count': len(text),
                'success': len(text.strip()) > 0
            }
            
        except Exception as e:
            return {
                'text': f"❌ Image extraction error: {str(e)}",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
    
    def _extract_docx(self) -> Dict[str, any]:
        """
        Extract text from DOCX file
        
        Returns:
            Dict: Extraction result with text and metadata
        """
        if not DOCX_AVAILABLE:
            return {
                'text': "❌ python-docx not available. Install: pip install python-docx",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
        
        try:
            try:
                doc = docx.Document(str(self.file_path))
            except Exception as e:
                error_msg = str(e).lower()
                if "password" in error_msg or "encrypted" in error_msg:
                    return {
                        'text': "❌ DOCX is password protected. python-docx has limited password support.",
                        'method': 'Error',
                        'file_name': self.file_path.name,
                        'success': False
                    }
                raise
            
            # Extract paragraphs
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Extract tables
            table_text = ""
            for table_num, table in enumerate(doc.tables, 1):
                table_text += f"\n\n--- Table {table_num} ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
            
            full_text = text + table_text
            
            return {
                'text': full_text.strip(),
                'method': 'python-docx',
                'file_name': self.file_path.name,
                'file_size': f"{self.file_path.stat().st_size / 1024:.2f}KB",
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'char_count': len(full_text),
                'success': True
            }
            
        except Exception as e:
            return {
                'text': f"❌ DOCX extraction error: {str(e)}",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
    
    def _extract_csv(self) -> Dict[str, any]:
        """
        Extract text from CSV file
        
        Returns:
            Dict: Extraction result with formatted table
        """
        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                rows = list(reader)
            
            # Format as markdown table
            text = "# CSV Data\n\n"
            
            if rows:
                # Headers
                headers = rows[0]
                text += "| " + " | ".join(headers) + " |\n"
                text += "| " + " | ".join(['---'] * len(headers)) + " |\n"
                
                # Data rows (limit to first 100)
                for row in rows[1:101]:
                    if len(row) == len(headers):
                        text += "| " + " | ".join(row) + " |\n"
                
                if len(rows) > 101:
                    text += f"\n*Showing 100 rows of {len(rows)-1} total*\n"
            
            return {
                'text': text,
                'method': 'CSV Reader',
                'file_name': self.file_path.name,
                'file_size': f"{self.file_path.stat().st_size / 1024:.2f}KB",
                'row_count': len(rows),
                'column_count': len(rows[0]) if rows else 0,
                'char_count': len(text),
                'success': True
            }
            
        except Exception as e:
            return {
                'text': f"❌ CSV extraction error: {str(e)}",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
    
    def _extract_excel(self) -> Dict[str, any]:
        """
        Extract text from Excel file
        
        Returns:
            Dict: Extraction result with all sheets data
        """
        if not EXCEL_AVAILABLE:
            return {
                'text': "❌ Excel libraries not available. Install: pip install openpyxl pandas",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
        
        try:
            try:
                excel_file = pd.ExcelFile(str(self.file_path), engine='openpyxl')
            except Exception as e:
                error_msg = str(e).lower()
                if "password" in error_msg:
                    return {
                        'text': "❌ Excel file is password protected.",
                        'method': 'Error',
                        'file_name': self.file_path.name,
                        'success': False
                    }
                raise
            
            text = f"# Excel Data - {self.file_path.name}\n\n"
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                text += f"## Sheet: {sheet_name}\n\n"
                if hasattr(df, 'to_markdown'):
                    text += df.to_markdown(index=False)
                else:
                    text += df.to_string()
                text += "\n\n"
            
            return {
                'text': text,
                'method': 'pandas + openpyxl',
                'file_name': self.file_path.name,
                'file_size': f"{self.file_path.stat().st_size / 1024:.2f}KB",
                'sheet_count': len(excel_file.sheet_names),
                'char_count': len(text),
                'success': True
            }
            
        except Exception as e:
            return {
                'text': f"❌ Excel extraction error: {str(e)}",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
    
    def _extract_text(self) -> Dict[str, any]:
        """
        Extract text from plain text file
        
        Returns:
            Dict: Extraction result with text and metadata
        """
        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            return {
                'text': text,
                'method': 'Plain Text',
                'file_name': self.file_path.name,
                'file_size': f"{self.file_path.stat().st_size / 1024:.2f}KB",
                'char_count': len(text),
                'line_count': len(text.split('\n')),
                'success': True
            }
            
        except Exception as e:
            return {
                'text': f"❌ Text extraction error: {str(e)}",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
    
    def _extract_generic(self) -> Dict[str, any]:
        """
        Generic extraction for unknown file types
        
        Returns:
            Dict: Extraction result 
        """
        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            return {
                'text': text,
                'method': 'Generic Text Reader',
                'file_name': self.file_path.name,
                'file_size': f"{self.file_path.stat().st_size / 1024:.2f}KB",
                'note': 'Unsupported file type - attempted generic text extraction',
                'success': len(text.strip()) > 0
            }
            
        except Exception as e:
            return {
                'text': f"❌ Unsupported file type: {self.extension}",
                'method': 'Error',
                'file_name': self.file_path.name,
                'success': False
            }
    
    def _estimate_pages(self, text: str) -> int:
        """
        Estimate page count from extracted text
        
        Args:
            text (str): Extracted text
            
        Returns:
            int: Estimated page count
        """
        chars_per_page = 3000
        return max(1, len(text) // chars_per_page)


def extract_from_file(file_path: str, output_file: Optional[str] = None, 
                     password: Optional[str] = None) -> Dict[str, any]:
    """
    Extract text from a document file with optional password
    
    Args:
        file_path (str): Path to the document
        output_file (str, optional): Path to save extracted text
        password (str, optional): Password for encrypted documents
    
    Returns:
        Dict: Dictionary with extraction results
    """
    extractor = DocumentExtractor(file_path, password=password)
    result = extractor.extract()
    
    # Save to file if requested and extraction successful
    if output_file and result['success']:
        output_path = Path(output_file)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result['text'])
        
        print(f"✓ Text saved to: {output_path}")
        result['output_file'] = str(output_path)
    
    return result


def batch_extract(directory: str, output_dir: str = "./extracted", 
                 password: Optional[str] = None):
    """
    Extract text from all documents in a directory
    
    Args:
        directory (str): Path to directory containing documents
        output_dir (str): Directory to save extracted texts
        password (str, optional): Password for encrypted documents
    """
    dir_path = Path(directory)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    supported_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.bmp', 
                           '.gif', '.docx', '.doc', '.csv', '.xlsx', '.xls', '.txt']
    files = [f for f in dir_path.iterdir() 
            if f.suffix.lower() in supported_extensions and f.is_file()]
    
    if not files:
        print(f"⚠ No supported documents found in {dir_path}")
        return
    
    print(f"\n⟳ Found {len(files)} documents to process\n")
    
    results = []
    successful = 0
    failed = 0
    
    for file in files:
        try:
            output_file = output_path / f"{file.stem}.txt"
            result = extract_from_file(str(file), str(output_file), password=password)
            results.append(result)
            
            if result['success']:
                successful += 1
                status = "✓"
            else:
                failed += 1
                status = "✗"
            
            char_count = result.get('char_count', 0)
            print(f"{status} {file.name} - {char_count} chars")
            
        except Exception as e:
            failed += 1
            print(f"✗ {file.name} - Error: {str(e)}")
    
    print(f"\n{'='*50}")
    print(f"✓ Processing complete!")
    print(f"  Total: {len(results)} files")
    print(f"  Successful: {successful}")
    print(f"  Failed: {failed}")
    print(f"  Output: {output_path.absolute()}")
    print(f"{'='*50}")


def print_dependencies():
    """Print dependency check status"""
    print("\n" + "="*50)
    print("DEPENDENCY CHECK")
    print("="*50)
    print(f"  PyMuPDF:     {'✓ Available' if PYMUPDF_AVAILABLE else '✗ Install: pip install PyMuPDF'}")
    print(f"  pypdf:       {'✓ Available' if PYPDF_AVAILABLE else '✗ Install: pip install pypdf'}")
    print(f"  pdfplumber:  {'✓ Available' if PDFPLUMBER_AVAILABLE else '✗ Install: pip install pdfplumber'}")
    print(f"  tabula:      {'✓ Available' if TABULA_AVAILABLE else '✗ Install: pip install tabula-py'}")
    print(f"  OCR:         {'✓ Available' if OCR_AVAILABLE else '✗ Install: pip install pytesseract pdf2image pillow'}")
    print(f"  DOCX:        {'✓ Available' if DOCX_AVAILABLE else '✗ Install: pip install python-docx'}")
    print(f"  Excel:       {'✓ Available' if EXCEL_AVAILABLE else '✗ Install: pip install openpyxl pandas'}")
    print("="*50 + "\n")


def main():
    """CLI interface with password support"""
    parser = argparse.ArgumentParser(
        description="🔓 Universal Document Text Extractor (Enhanced with Table Support)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
╔════════════════════════════════════════════════════════════╗
║                    USAGE EXAMPLES                          ║
╚════════════════════════════════════════════════════════════╝

1. Simple extraction:
   python extractor.py document.pdf

2. Save to file:
   python extractor.py statement.pdf -o output.txt

3. Password-protected PDF:
   python extractor.py locked.pdf --password "mypassword" -o unlocked.txt
   
4. Short form with password:
   python extractor.py locked.pdf -p "pass123" -o output.txt

5. Extract with tables (auto-detected):
   python extractor.py bank_statement.pdf -o statement.txt

6. Batch processing:
   python extractor.py --batch ./documents --output-dir ./extracted

7. Batch with password:
   python extractor.py --batch ./documents --password "yourpass"

8. Show dependencies:
   python extractor.py --check

9. Check if document is password protected:
   python extractor.py document.pdf --check-protected
        """
    )
    
    parser.add_argument('file', nargs='?', help='File to extract text from')
    parser.add_argument('--check-protected', action='store_true', 
                       help='Check if the document is password protected')
    parser.add_argument('-o', '--output', help='Output file path')
    parser.add_argument('-p', '--password', help='Password for encrypted/protected documents')
    parser.add_argument('--batch', help='Process all files in directory')
    parser.add_argument('--output-dir', default='./extracted', 
                       help='Output directory for batch processing (default: ./extracted)')
    parser.add_argument('--check', action='store_true', help='Check dependencies')
    
    args = parser.parse_args()
    
    # Print dependencies if requested
    if args.check:
        print_dependencies()
        return
    
    # Print dependencies status
    print_dependencies()
    
    if args.batch:
        # Batch processing
        batch_extract(args.batch, args.output_dir, password=args.password)
    elif args.check_protected:
        if not args.file:
            print("⚠ Please provide a file to check protection status.")
            return
        extractor = DocumentExtractor(args.file, password=args.password)
        if extractor.extension == '.pdf':
            protected = False
            if PYMUPDF_AVAILABLE:
                try:
                    doc = fitz.open(str(extractor.file_path))
                    protected = doc.is_encrypted
                    doc.close()
                except Exception as e:
                    print(f"Error checking PDF protection: {e}")
            elif PYPDF_AVAILABLE:
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(str(extractor.file_path))
                    protected = reader.is_encrypted
                except Exception as e:
                    print(f"Error checking PDF protection: {e}")
            print(f"{'✓ Protected' if protected else '✗ Not protected'}: {extractor.file_path.name}")
        else:
            print(f"⚠ Protection check not supported for: {extractor.extension}")
        return
    elif args.file:
        # Single file extraction
        result = extract_from_file(args.file, args.output, password=args.password)
        
        print("\n" + "="*50)
        if result['success']:
            print("✓ EXTRACTION SUCCESSFUL")
            print("="*50)
            print(f"  File: {result['file_name']}")
            print(f"  Method: {result['method']}")
            print(f"  Size: {result['file_size']}")
            print(f"  Characters: {result['char_count']:,}")
            print(f"  Pages: {result.get('page_count', 'N/A')}")
            
            if result.get('has_tables'):
                print(f"  Tables: ✓ Detected and extracted")
            
            if result.get('output_file'):
                print(f"  Output: {result['output_file']}")
            
            # Show preview if not saved
            if not args.output:
                print("\n" + "-"*50)
                print("PREVIEW:")
                print("-"*50)
                preview_text = result['text']
                print(preview_text)

        else:
            print("✗ EXTRACTION FAILED")
            print("="*50)
            print(f"  File: {result['file_name']}")
            print(f"  Error: {result['text']}")
        
        print("="*50 + "\n")
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
